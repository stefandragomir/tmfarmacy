
import os
import sys
import random
import base64

try:
    _path = os.path.split(os.path.split(os.path.abspath(__file__))[0])[0]
    sys.path.append(_path)
except:
    pass

_path = os.path.split(os.path.split(os.path.abspath("__file__"))[0])[0]
sys.path.append(_path)

from docx                         import Document
from docx.enum.text               import WD_ALIGN_PARAGRAPH
from PyQt5.QtCore                 import *
from PyQt5.QtGui                  import *
from PyQt5.QtWidgets              import * 
from pharm_widgets.pharm_widgets  import *
from pharm_widgets.pharm_css      import *
from pharm_icons.pharm_icons      import Pharm_Icon
from pharm_icons.pharm_icons      import Pharm_Pixmap
from datetime                     import datetime
from pharm_db.pharm_db            import PHARM_DB
from docx                         import Document
from docx.enum.text               import WD_BREAK
from datetime                     import datetime
from functools                    import partial
from docx.oxml                    import OxmlElement
from docx.oxml.ns                 import qn
from io                           import StringIO
from io                           import BytesIO
from pprint                       import pprint

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""

PHARM_NUMBER_OF_TEST_QUESTIONS = 10
PHARM_MIN_CORECT_QUESTIONS     = 8
PHARM_ANSWER_OPTIONS           = ["a","b","c","d","e","f","g","h"]
PHARM_QUESTIONS_HEADING        = """
INTREBARI EXAMEN ACORDARE/PRELUNGIRE
LICENŢĂ PILOT AERONAVE ULTRAUŞOARE
CLASA PARAPANTĂ
""" 
PHARM_TEST_HEADING        = """
CHESTIONAR EXAMEN ACORDARE/PRELUNGIRE
LICENŢĂ PILOT AERONAVE ULTRAUŞOARE
CLASA PARAPANTĂ
""" 

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_UI(QMainWindow):

    def __init__(self,db):

        QMainWindow.__init__(self)

        self.stack_index  = 0
        self.db           = db   
        self.dekstops     = []   

        self.draw_gui() 

    def draw_gui(self):

        self.setWindowTitle("Admitere Farmacie 1.0.0")
        self.setWindowIcon(Pharm_Icon("pharm"))
        self.setMinimumSize(1300, 800)       
        self.setMinimumHeight(500)

        self.setStyleSheet(PHARM_CSS)

        self.wdg_central = QWidget()
        
        #selection tree
        self.tree = Pharm_WDG_Tree()
        self.tree.setFixedWidth(280)
        self.populate_tree()
        self.tree.currentItemChanged.connect(self.tree_select)

        self.toolbar_layout = QHBoxLayout()

        self.bt_generate_all = Pharm_WDG_Button("Genereaza Teste",           Pharm_Icon("generate_test"), Pharm_Icon("generate_test"), "#606060")
        self.toolbar_layout.addWidget(self.bt_generate_all)

        self.bt_generate_all.clicked.connect(self.clbk_bt_gen_all)

        #left tree layout area
        self.tree_area = QVBoxLayout()   
        self.tree_area.addLayout(self.toolbar_layout) 
        self.tree_area.addWidget(self.tree)   

        #context stack
        self.context = Pharm_WDG_Stack()
        self.populate_stack()

        #Layout       
        self.top_layout = QHBoxLayout()
        self.top_layout.addLayout(self.tree_area)
        self.top_layout.addWidget(self.context)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.top_layout)

        self.wdg_central.setLayout(self.main_layout)  

        self.setCentralWidget(self.wdg_central)
        self.activateWindow() 

    def showEvent(self, event):    

        pass
        
    def tree_select(self, current, previous):

        self.context.setCurrentIndex(0) 

        _item_cfg = str(current.data(1,Qt.UserRole))

        #set context widget based on coresponding state     
        self.context.setCurrentIndex(int(_item_cfg))  

    def populate_tree(self):

        self.tree.clear()

        self.stack_index = 1 

        _parent = self.tree.invisibleRootItem()

        for _category in self.db:

            _item = QTreeWidgetItem(_parent)
            _item.setData(0, Qt.EditRole, _category.name)             

            _serial_cfg = self.stack_index                    
            _item.setData(1, Qt.UserRole, _serial_cfg)

            _item.setIcon(0,Pharm_Icon("test"))

            self.stack_index += 1

    def populate_stack(self):

        self.context.insertWidget(0, QWidget())

        for _idx in range(len(self.db)):

            _desktop = Pharm_WDG_Desktop(self.db[_idx])

            self.dekstops.append(_desktop)

            self.context.insertWidget(_idx + 1, _desktop) 

        self.context.setCurrentIndex(0)  

    def clbk_bt_gen_all(self,state):

        _paths = []

        _timestamp = str(datetime.now())

        for _desktop in self.dekstops:

            _path = _desktop.clbk_bt_gen(state=True,timestamp=_timestamp)

            _paths.append(_path)

        _path_txt = ""

        for _path in _paths:

            _path_txt += "   %s \n" % (os.path.split(_path)[1],)

        _msg = QMessageBox()
        _msg.setWindowIcon(Pharm_Icon("parag"))
        _msg.setIcon(QMessageBox.Information)
        _msg.setText("Generat chestionare pentru test:\n%s"  % (_path_txt,))
        _msg.setWindowTitle("Generat Chestionar Test " )
        _msg.exec_()

        os.startfile(os.path.split(_path)[0])

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_Model_Test(object):

    def __init__(self):

        self.questions = []

    def clear(self):

        for _question in self.questions:

            for _answer in _question.answers:

                _answer.is_selected = False

    def get_result(self):

        _total    = len(self.questions)
        _corect   = 0
        _incorect = 0

        for _question in self.questions:

            _user_scor = 0
            _nrm_scor  = 0

            for _answer in _question.answers:

                if _answer.corect and  _answer.selected:
                    _user_scor += 1

                if _answer.corect:
                    _nrm_scor += 1

            if _user_scor == _nrm_scor:
                _corect += 1
            else:
                _incorect += 1

        return _total,_corect,_incorect

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Desktop(QWidget):

    def __init__(self,category):

        QWidget.__init__(self)

        self.category = category

        self.draw_gui()

    def draw_gui(self):

        self.bt_test  = Pharm_WDG_Button("Test",              Pharm_Icon("test_normal"),   Pharm_Icon("test_hover"),  "#606060")
        self.bt_learn = Pharm_WDG_Button("Invata",            Pharm_Icon("learn_normal"),  Pharm_Icon("learn_hover"), "#606060")
        self.bt_gen   = Pharm_WDG_Button("Genereaza Test",    Pharm_Icon("generate_test"), Pharm_Icon("generate_test"), "#606060")
        self.bt_exp   = Pharm_WDG_Button("Exporta Intrebari", Pharm_Icon("generate_doc"), Pharm_Icon("generate_doc"), "#606060")

        self.bt_test.setIconSize(QSize(100,100))
        self.bt_learn.setIconSize(QSize(100,100))
        self.bt_gen.setIconSize(QSize(100,100))
        self.bt_exp.setIconSize(QSize(100,100))

        self.bt_test.clicked.connect(self.clbk_bt_test)
        self.bt_learn.clicked.connect(self.clbk_bt_learn)
        self.bt_gen.clicked.connect(self.clbk_bt_gen)
        self.bt_exp.clicked.connect(self.clbk_bt_exp)

        self.wdg_test = Pharm_WDG_Desktop_Test(self,self.category)

        self.bt_layout = QHBoxLayout()
        self.bt_layout.addWidget(self.bt_learn)
        self.bt_layout.addWidget(self.bt_test)
        self.bt_layout.addWidget(self.bt_gen)
        self.bt_layout.addWidget(self.bt_exp)

        self.main_layout = QVBoxLayout()
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_test)

        self.wdg_test.hide()

        self.setLayout(self.main_layout)  

    def get_docs_path(self):

        _path = os.path.abspath(__file__)
        _path = os.path.split(_path)[0]
        _path = os.path.join(_path,"docs")

        if not os.path.exists(_path):

            os.mkdir(_path)

        return _path

    def get_generated_doc_path(self,name):

        _path = self.get_docs_path()

        _timestamp = str(datetime.now())

        _file_name = "intrebari_%s_%s.docx" % (name,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

        _path = os.path.join(_path,_file_name)

        return _path,_timestamp

    def get_generated_test_path(self,name,timestamp=None):

        _path = self.get_docs_path()

        if timestamp == None:
            _timestamp = str(datetime.now())
        else:
            _timestamp = timestamp
            _path = os.path.join(_path,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

            if not os.path.exists(_path):

                os.mkdir(_path)

        _file_name = "test_%s_%s.docx" % (name,_timestamp.replace(":","_").replace(" ","_").replace("/","_").replace("-","_").replace(".","_"))

        _path = os.path.join(_path,_file_name)

        return _path,_timestamp

    def clbk_bt_test(self,state):

        self.wdg_test.show()

        self.bt_test.hide()

        self.bt_learn.hide()

        self.bt_gen.hide()

        self.bt_exp.hide()

        self.wdg_test.start("test")

    def clbk_bt_learn(self,state):

        self.wdg_test.show()

        self.bt_test.hide()

        self.bt_learn.hide()

        self.bt_gen.hide()

        self.bt_exp.hide()

        self.wdg_test.start("learn")

    def clbk_bt_gen(self,state,timestamp=None):

        _questions_indexes = random.sample(range(len(self.category.questions)), PHARM_NUMBER_OF_TEST_QUESTIONS)

        _path, _timestamp = self.get_generated_test_path(self.category.name,timestamp)

        _document = Document()

        _paragraph_format             = _document.styles['Normal'].paragraph_format
        _paragraph_format.space_after = 3

        self.__generate_top_table(_document)

        _paragraph = _document.add_paragraph()
        _paragraph.add_run("%s\n%s" % (PHARM_TEST_HEADING,self.category.name.upper())).bold = True
        _paragraph.alignment = WD_ALIGN_PHARMRAPH.CENTER
        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        _question_count = 0

        _table_questions = []

        for _index in range(len(self.category.questions)):

            if _index in _questions_indexes:

                _question = self.category.questions[_index]

                _table_questions.append([_index + 1,_question])

                _question_text = "%s. %s" % (_question_count + 1,_question.text)

                _question_count += 1

                _paragraph = _document.add_paragraph(_question_text)

                if _question.image != None:

                    _bytes = base64.b64decode(_question.image)

                    _str = BytesIO(_bytes)

                    _document.add_picture(_str)

                _answer_count = 0

                for _answer in _question.answers:

                    _answer_text = "    %s) %s" % (PHARM_ANSWER_OPTIONS[_answer_count],_answer.text)

                    _answer_count += 1

                    _paragraph = _document.add_paragraph(_answer_text)

        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        self.__generate_bottom_table(_document,_table_questions)

        _document.add_paragraph("Test generat la data: %s" % (_timestamp,))

        _document.save(_path)

        if timestamp == None:

            _msg = QMessageBox()
            _msg.setWindowIcon(Pharm_Icon("parag"))
            _msg.setIcon(QMessageBox.Information)
            _msg.setText("Generat chestionar test in fisierul\n%s" % (os.path.split(_path)[1],))
            _msg.setWindowTitle("Generat Chestionar Test %s" % (self.category.name.upper(),))
            _msg.exec_()

        
            os.startfile(_path)

        return _path

    def __generate_top_table(self,document):

            _table = document.add_table(rows=2, cols=3)

            _table.cell(0, 0).text = "NUME SI PRENUME "
            self.set_cell_border(
                                    _table.cell(0, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 1).text = "SEMNATURA"
            self.set_cell_border(
                                    _table.cell(0, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 2).text = "DATA"
            self.set_cell_border(
                                    _table.cell(0, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            self.set_cell_border(
                                    _table.cell(1, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

    def __generate_bottom_table(self,document,questions):

            document.add_page_break()

            _paragraph = document.add_paragraph()
            _paragraph.add_run("%s\n%s" % (PHARM_TEST_HEADING,self.category.name.upper())).bold = True
            _paragraph.alignment = WD_ALIGN_PHARMRAPH.CENTER

            _table = document.add_table(rows=11, cols=4)

            _table.cell(0, 0).text = "Numar Intrebare Test"
            self.set_cell_border(
                                    _table.cell(0, 0),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 1).text = "Numar Intrebare Documentatie"
            self.set_cell_border(
                                    _table.cell(0, 1),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 2).text = "Raspuns Corect"
            self.set_cell_border(
                                    _table.cell(0, 2),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            _table.cell(0, 3).text = "Raspuns Examen"
            self.set_cell_border(
                                    _table.cell(0, 3),
                                    top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                    end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _table.cell(_index + 1, 0).text = str(_index + 1)  

                self.set_cell_border(
                                        _table.cell(_index + 1, 0),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _table.cell(_index + 1, 1).text = str(questions[_index][0])  

                self.set_cell_border(
                                        _table.cell(_index + 1, 1),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                _answer_text = ""

                _answer_count = 0

                for _answer in questions[_index][1].answers:

                    if _answer.corect:

                        _answer_text += "%s, " % (PHARM_ANSWER_OPTIONS[_answer_count],)

                    _answer_count += 1


                _table.cell(_index + 1, 2).text = _answer_text

                self.set_cell_border(
                                        _table.cell(_index + 1, 2),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

            for _index in range(len(questions)):

                self.set_cell_border(
                                        _table.cell(_index + 1, 3),
                                        top={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        bottom={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        start={"sz": 5, "val": "single", "color": "#000000", "space": "0"},
                                        end={"sz": 5, "val": "single", "color": "#000000", "space": "0"})

    def set_cell_border(sefl, cell, **kwargs):
        """
        Set cell`s border
        Usage:

        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def clbk_bt_exp(self,state):

        _path, _timestamp = self.get_generated_doc_path(self.category.name)

        _document = Document()

        _paragraph = _document.add_paragraph()
        _paragraph.add_run("%s\n%s" % (PHARM_QUESTIONS_HEADING,self.category.name.upper())).bold = True
        _paragraph.alignment = WD_ALIGN_PHARMRAPH.CENTER
        _document.add_paragraph("")
        _document.add_paragraph("")

        _question_count = 1

        for _question in self.category.questions:

            _question_text = "%s. %s" % (_question_count,_question.text)

            _question_count += 1

            _paragraph = _document.add_paragraph(_question_text)

            if _question.image != None:

                _bytes = base64.b64decode(_question.image)

                _str = BytesIO(_bytes)

                _document.add_picture(_str)

            _answer_count = 0


            for _answer in _question.answers:

                _answer_text = "    %s) %s" % (PHARM_ANSWER_OPTIONS[_answer_count],_answer.text)

                _answer_count += 1

                _paragraph = _document.add_paragraph()
                _run       = _paragraph.add_run(_answer_text)
                _run.bold  = _answer.corect

            _document.add_paragraph("")

        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("")
        _document.add_paragraph("Generat la data: %s" % (_timestamp,))

        _document.save(_path)

        _msg = QMessageBox()
        _msg.setWindowIcon(Pharm_Icon("parag"))
        _msg.setIcon(QMessageBox.Information)
        _msg.setText("Generat intrebari in fisierul\n%s" % (os.path.split(_path)[1],))
        _msg.setWindowTitle("Generat Intrebari %s" % (self.category.name.upper(),))
        _msg.exec_()

        os.startfile(_path)

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Desktop_Test(QWidget):

    def __init__(self,parent,category):

        QWidget.__init__(self)

        self.category        = category
        self.question_number = 0
        self.parent          = parent
        self.test_type       = ""
        self.test_learn      = Pharm_Model_Test()
        self.test_exam       = Pharm_Model_Test()

        self.test_learn.questions = self.category.questions

        self.draw_gui()

    def draw_gui(self):

        self.bt_next   = Pharm_WDG_Small_Button( Pharm_Icon("next_normal"),       Pharm_Icon("next_hover"),     self.clbk_next , "Intrebare urmatoare")
        self.bt_prev   = Pharm_WDG_Small_Button( Pharm_Icon("previous_normal"),   Pharm_Icon("previous_hover"), self.clbk_prev, "Intrebarea precedenta")
        self.bt_result = Pharm_WDG_Small_Button( Pharm_Icon("results_normal"),    Pharm_Icon("result_hover"),   self.clbk_result, "Rezultat")
        self.bt_close  = Pharm_WDG_Small_Button( Pharm_Icon("close_normal"),      Pharm_Icon("close_hover"),    self.clbk_close, "Inchide")

        self.bt_next.setIconSize(QSize(50,50))
        self.bt_prev.setIconSize(QSize(50,50))
        self.bt_result.setIconSize(QSize(50,50))
        self.bt_close.setIconSize(QSize(50,50))

        self.lbl_status = Pharm_WDG_Label()
        self.lbl_result = Pharm_WDG_Label()

        self.bt_layout = QHBoxLayout()
        
        self.bt_layout.addWidget(self.bt_close)
        self.bt_layout.addWidget(self.bt_prev)
        self.bt_layout.addWidget(self.bt_result)
        self.bt_layout.addWidget(self.bt_next)

        self.wdg_question = Pharm_WDG_Question()

        self.main_layout = QVBoxLayout()
        
        self.main_layout.addLayout(self.bt_layout)
        self.main_layout.addWidget(self.wdg_question)
        self.main_layout.addWidget(self.lbl_status)
        self.main_layout.addWidget(self.lbl_result)

        self.setLayout(self.main_layout) 

        self.lbl_result.hide()

        self.bt_next.hide()
        self.bt_prev.hide()
        self.bt_close.hide()
        self.bt_result.hide()

    def start(self,test_type):

        self.test_type = test_type

        self.question_number = 0

        self.wdg_question.show()

        if self.test_type == "learn":            
            self.test_learn.clear()
            self.wdg_question.populate(self.test_learn.questions[self.question_number])
            self.bt_next.show()
            self.bt_close.show()
            self.bt_result.show()
            self.bt_prev.show()
        else:            
            self.test_exam.clear()
            self.bt_close.show()
            self.bt_next.show()
            self.bt_prev.show()
            self.get_test_questions()
            self.wdg_question.populate(self.test_exam.questions[self.question_number])

        self.lbl_result.hide()
        self.set_status()

    def get_test_questions(self):

        global PHARM_NUMBER_OF_TEST_QUESTIONS

        if PHARM_NUMBER_OF_TEST_QUESTIONS < len(self.test_learn.questions):

            _questions_indexes = random.sample(range(len(self.test_learn.questions)), PHARM_NUMBER_OF_TEST_QUESTIONS)

            self.test_exam.questions = [self.test_learn.questions[_index] for _index in _questions_indexes]
        else:
            self.test_exam.questions = self.test_learn.questions

    def clbk_next(self,state):

        if not self.is_end(self.question_number):

            self.question_number += 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number])
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number])

            self.set_status()
        else:
            pass

        if self.test_type != "learn":
            if self.question_number == len(self.test_exam.questions) - 1:
                self.bt_result.show()
            else:
                self.bt_result.hide()

    def clbk_prev(self,state):

        if self.test_type != "learn":

            self.bt_result.hide()

        if not self.is_begining(self.question_number):

            self.question_number -= 1

            if self.test_type == "learn":
                self.wdg_question.populate(self.test_learn.questions[self.question_number])
            else:
                self.wdg_question.populate(self.test_exam.questions[self.question_number])

            self.set_status()

        else:
            pass

    def clbk_close(self,state):

        self.parent.bt_test.show()
        self.parent.bt_learn.show()
        self.parent.bt_gen.show()
        self.parent.bt_exp.show()
        self.parent.wdg_test.hide()

    def clbk_result(self,state):

        _corect   = 0
        _incorect = 0
        _total    = 0

        if self.test_type == "learn":

            for _index in range(len(self.test_learn.questions[self.question_number].answers)):

                if self.test_learn.questions[self.question_number].answers[_index].corect:
                    self.wdg_question.rd_answers[_index].setStyleSheet("QCheckBox { color: green }")
                else:
                    self.wdg_question.rd_answers[_index].setStyleSheet("QCheckBox { color: red }")
        else:
            _total,_corect,_incorect = self.test_exam.get_result()

            self.lbl_status.setText("Intrebari[%s] Corecte[%s] Incorecte[%s]" % (_total,_corect,_incorect))

            self.bt_next.hide()
            self.bt_prev.hide()
            self.bt_result.hide()
            self.bt_close.show()
            self.wdg_question.hide()

            self.lbl_result.show()

            if _corect >= PHARM_MIN_CORECT_QUESTIONS:
                self.lbl_result.setText("ADMIS")
                self.lbl_result.setStyleSheet("QLabel { background-color : #14c941; font: 18pt; color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)
            else:
                self.lbl_result.setText("PICAT")
                self.lbl_result.setStyleSheet("QLabel { background-color : #ba2012; font: 18pt;  color: #ffffff}")
                self.lbl_result.setAlignment(Qt.AlignCenter)

    def set_status(self):

        if self.test_type == "learn":
            self.lbl_status.setText("Intrebarea[%s/%s] " % (self.question_number + 1,len(self.test_learn.questions)))
        else:
            self.lbl_status.setText("Intrebarea[%s/%s]" % (self.question_number + 1,len(self.test_exam.questions)))

    def is_end(self,question_number):

        _state = False

        if self.test_type == "learn":
            _state = (question_number + 1) >= len(self.test_learn.questions)
        else:
            _state = (question_number + 1) >= len(self.test_exam.questions)

        return _state

    def is_begining(self,question_number):

        return question_number <= 0

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm_WDG_Question(QWidget):

    def __init__(self):

        QWidget.__init__(self)

        self.draw_gui()

        self.question = None

    def draw_gui(self):

        self.lbl_question = Pharm_WDG_Label()
        self.lbl_question.setWordWrap(True)
        self.lbl_image    = Pharm_WDG_Label()
        self.rd_answers   = []

        for _index in range(5):
            self.rd_answers.append(Pharm_WDG_CheckBox(""))
            self.rd_answers[-1].checkbox.stateChanged.connect(partial(self.clbk_answer,_index))

        self.main_layout = QVBoxLayout()
        self.main_layout.addWidget(self.lbl_question)
        self.main_layout.addWidget(self.lbl_image)

        for _index in range(5):
            self.main_layout.addWidget(self.rd_answers[_index])

        self.setLayout(self.main_layout)

        self.lbl_question.hide()
        self.lbl_image.hide()

        for _index in range(5):

            self.rd_answers[_index].hide()

    def populate(self,question):

        self.lbl_image.hide()

        for _index in range(5):

            self.rd_answers[_index].hide()

        self.question = question

        self.lbl_question.show()

        if question.image != None:

            self.lbl_image.show()

            _barray      = QByteArray()
            _barray_data = _barray.fromBase64(question.image.encode("utf-8"))

            #create pixmap from base64 data
            _pixmap = QPixmap ()
            _pixmap.loadFromData(_barray_data, "PNG")

            _w = 600
            _h = 600

            self.lbl_image.setPixmap(_pixmap.scaled(_w,_h,Qt.KeepAspectRatio))
            self.lbl_image.setMask(_pixmap.mask())


        for _index in range(len(self.question.answers)):

            self.rd_answers[_index].show()
            self.rd_answers[_index].setStyleSheet("QCheckBox { color: #b1b1b1 }")
            self.rd_answers[_index].label.setText(self.question.answers[_index].text)

            if self.question.answers[_index].selected:
                self.rd_answers[_index].checkbox.setCheckState(Qt.Checked)
            else:
                self.rd_answers[_index].checkbox.setCheckState(Qt.Unchecked)

        self.lbl_question.setText(self.question.text)

    def clbk_answer(self,index,state):
        
        self.question.answers[index].selected = state == Qt.Checked

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
class Pharm(object):

    def __init__(self):

        global PHARM_DB

        self.app = QApplication(sys.argv)  

        _ui  = Pharm_UI(PHARM_DB)   

        _ui.show()

        sys.exit(self.app.exec_())

"""*************************************************************************************************
****************************************************************************************************
*************************************************************************************************"""
if __name__ == "__main__":

    Pharm()

